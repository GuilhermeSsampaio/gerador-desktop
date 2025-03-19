from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches


def create_document():
    """Cria e retorna um novo documento com configurações padrão."""
    return Document()

def add_title(doc, title):
    """Adiciona um título centralizado e em negrito."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Cria o run com o texto e aplica negrito
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    return paragraph

def add_subtitle(doc, subtitle):
    """Adiciona um subtítulo em negrito e alinhado à esquerda."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Cria o run com o texto e aplica negrito
    run = paragraph.add_run(subtitle)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    return paragraph

def add_paragraph_text(doc, text, bold=False, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """Adiciona um parágrafo com formatação específica."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    
    # Cria o run com o texto
    run = paragraph.add_run(text)
    
    # Aplica negrito se solicitado
    if bold:
        run.bold = True
    
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    return paragraph

def add_empty_paragraph(doc):
    """Adiciona um parágrafo vazio."""
    return doc.add_paragraph()

def add_page_break(doc):
    """Adiciona uma quebra de página se o último parágrafo não estiver vazio."""
    if doc.paragraphs and doc.paragraphs[-1].text.strip():
        doc.add_page_break()

def save_document(doc, file_path):
    """Salva o documento no caminho especificado."""
    doc.save(file_path)
    

from docx.shared import Inches, Pt
from PIL import Image
import os

def add_image(doc, image_path, width=4, height=None, show_caption=False, add_page_break_after=False):
    """
    Adiciona uma imagem ao documento com o tamanho especificado.
    Tenta diferentes abordagens se a primeira falhar.
    
    Args:
        doc: O documento docx
        image_path: Caminho completo para a imagem
        width: Largura da imagem em polegadas (padrão: 4)
        height: Altura da imagem em polegadas (opcional)
        show_caption: Se True, adiciona o nome do arquivo como legenda (padrão: False)
        add_page_break_after: Se True, adiciona uma quebra de página após a imagem (padrão: False)
    """
    try:
        # Verificar se o arquivo existe
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Arquivo de imagem não encontrado: {image_path}")
            
        # Primeiro tenta o método direto
        if height:
            doc.add_picture(image_path, width=Inches(width), height=Inches(height))
        else:
            # Se a altura não for especificada, mantém a proporção original
            doc.add_picture(image_path, width=Inches(width))
        
        # Adiciona o nome do arquivo como legenda apenas se show_caption for True
        if show_caption:
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1  # Centralizado
            file_name = os.path.basename(image_path)
            run = paragraph.add_run(f"Imagem: {file_name}")
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
        
        # Adiciona uma quebra de página após a imagem, se solicitado
        if add_page_break_after:
            doc.add_page_break()
        
        return True
        
    except Exception as e:
        try:
            # Segunda tentativa: usar PIL para abrir e validar a imagem primeiro
            img = Image.open(image_path)
            # Criar um arquivo temporário com formato explicitamente suportado
            temp_path = image_path + ".temp.png"
            img.save(temp_path)
            
            # Tentar adicionar com o arquivo temporário
            if height:
                doc.add_picture(temp_path, width=Inches(width), height=Inches(height))
            else:
                doc.add_picture(temp_path, width=Inches(width))
                
            # Remover arquivo temporário
            os.remove(temp_path)
            
            # Adiciona o nome do arquivo como legenda apenas se show_caption for True
            if show_caption:
                paragraph = doc.add_paragraph()
                paragraph.alignment = 1  # Centralizado
                file_name = os.path.basename(image_path)
                run = paragraph.add_run(f"Imagem: {file_name}")
                run.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
            
            # Adiciona uma quebra de página após a imagem, se solicitado
            if add_page_break_after:
                doc.add_page_break()
            
            return True
            
        except Exception as e2:
            # Em caso de falha nas duas tentativas, adicionar um texto informando o problema
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"Erro ao adicionar imagem {os.path.basename(image_path)}:")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.bold = True
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"Caminho: {image_path}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"Erro 1: {str(e)}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"Erro 2: {str(e2)}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            
            return False

# Função para verificar se o negrito está funcionando (teste)
def test_bold_functionality(output_path="test_bold.docx"):
    doc = create_document()
    
    # Adiciona título em negrito
    add_title(doc, "Este é um título em negrito")
    
    # Adiciona subtítulo em negrito
    add_subtitle(doc, "Este é um subtítulo em negrito")
    
    # Adiciona texto normal
    add_paragraph_text(doc, "Este é um texto normal")
    
    # Adiciona texto em negrito
    add_paragraph_text(doc, "Este texto deve estar em negrito", bold=True)
    
    # Salva para testar
    save_document(doc, output_path)
    print(f"Documento de teste salvo como {output_path}")