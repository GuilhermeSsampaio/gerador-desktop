import argparse
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import re

# Verifica se o módulo docx2pdf está disponível
try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("AVISO: O módulo 'docx2pdf' não está instalado. Apenas o documento DOCX será atualizado.")

def parse_args():
    """Processa os argumentos da linha de comando"""
    parser = argparse.ArgumentParser(description="Atualizador de arquivo de especificação")
    parser.add_argument("--title", required=True, help="Novo título para o documento")
    parser.add_argument("--date", required=True, help="Nova data para o documento (formato: DD/MM/YYYY)")
    parser.add_argument("--file", required=True, help="Caminho para o arquivo de especificação")
    return parser.parse_args()

def update_espec_file(title, date, file_path):
    """Atualiza o título e a data no arquivo de especificação"""
    if not os.path.exists(file_path):
        print(f"Erro: Arquivo não encontrado: {file_path}")
        return False
    
    try:
        # Abre o documento
        doc = Document(file_path)
        
        # Atualiza o título (primeiro parágrafo ou busca por um título específico)
        title_updated = False
        for i, paragraph in enumerate(doc.paragraphs):
            # Procura parágrafos que pareçam ser um título
            if i < 3 and paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                # Este é provavelmente o título
                paragraph.clear()
                run = paragraph.add_run(title)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                title_updated = True
                break
            # Outra opção: procurar por texto que contenha "Linguagem de Programação"
            elif "Linguagem de Programação" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run(title)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                title_updated = True
                break
        
        # Se não encontrou um título para atualizar, avisa
        if not title_updated:
            print("Aviso: Não foi possível identificar o parágrafo do título para atualização.")
        
        # Atualiza a data e assinatura (geralmente nos últimos parágrafos)
        date_updated = False
        for i in range(len(doc.paragraphs) - 1, max(0, len(doc.paragraphs) - 5), -1):
            paragraph = doc.paragraphs[i]
            # Procura parágrafos que contenham "Dourados" e uma data
            if "Dourados" in paragraph.text or re.search(r'\d{2}/\d{2}/\d{4}', paragraph.text):
                paragraph.clear()
                run = paragraph.add_run(f"Dourados, {date} -- (Assinatura)")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                date_updated = True
                break
        
        # Se não encontrou uma data para atualizar, avisa
        if not date_updated:
            print("Aviso: Não foi possível identificar o parágrafo da data para atualização.")
        
        # Salva o documento
        doc.save(file_path)
        print(f"Documento atualizado: {file_path}")
        
        # Converte para PDF se o arquivo original for DOCX
        if file_path.endswith('.docx') and PDF_AVAILABLE:
            try:
                pdf_path = file_path.replace('.docx', '.pdf')
                convert(file_path, pdf_path)
                print(f"PDF atualizado: {pdf_path}")
            except Exception as e:
                print(f"Erro ao converter especificação para PDF: {str(e)}")
        
        return True
    
    except Exception as e:
        print(f"Erro ao atualizar o arquivo de especificação: {str(e)}")
        return False

if __name__ == "__main__":
    args = parse_args()
    update_espec_file(args.title, args.date, args.file)
