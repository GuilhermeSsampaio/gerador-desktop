import os
import sys
import argparse
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Verifica se o módulo docx2pdf está disponível
try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("AVISO: O módulo 'docx2pdf' não está instalado. Apenas o documento DOCX será gerado.")

def parse_args():
    """Processa os argumentos da linha de comando"""
    parser = argparse.ArgumentParser(description="Gerador de documento de screenshots")
    parser.add_argument("--name", required=True, help="Nome completo do aluno")
    parser.add_argument("--entrega", required=True, help="Número da entrega")
    parser.add_argument("--screenshots-list", required=True, help="Arquivo com lista de caminhos para screenshots")
    parser.add_argument("--output", required=True, help="Nome do arquivo de saída (sem extensão)")
    return parser.parse_args()

def create_document():
    """Cria um novo documento Word"""
    return Document()

def add_title(doc, title):
    """Adiciona um título centralizado e em negrito"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Cria o run com o texto e aplica negrito
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = 16  # Tamanho equivalente a Pt(12)
    
    return paragraph

def add_empty_paragraph(doc):
    """Adiciona um parágrafo vazio"""
    return doc.add_paragraph()

def add_image(doc, image_path, caption=None):
    """Adiciona uma imagem ao documento com uma legenda opcional"""
    try:
        # Adiciona a imagem com largura fixa (ajustável)
        doc.add_picture(image_path, width=Inches(6.0))
        
        # Se houver legenda, adiciona-a
        if caption:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(caption)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = 12
        
        # Adiciona quebra de página após a imagem
        doc.add_page_break()
        return True
    except Exception as e:
        print(f"Erro ao adicionar imagem {image_path}: {e}")
        return False

def generate_screenshots_document(name, entrega, screenshots_list_file, output_path):
    """Gera o documento com screenshots"""
    doc = create_document()
    
    # Adiciona título
    title = f"Linguagem de Programação III - Entrega {entrega} - {name}"
    add_title(doc, title)
    add_empty_paragraph(doc)
    
    # Lê a lista de screenshots
    with open(screenshots_list_file, 'r') as f:
        screenshot_paths = [line.strip() for line in f if line.strip()]
    
    # Verifica se há screenshots para adicionar
    if not screenshot_paths:
        print("Nenhum screenshot encontrado para adicionar ao documento.")
        return False
    
    # Adiciona cada screenshot como uma página separada
    for i, screenshot_path in enumerate(screenshot_paths):
        if os.path.exists(screenshot_path):
            # Gera uma legenda baseada no nome do arquivo
            filename = os.path.basename(screenshot_path)
            base_name = os.path.splitext(filename)[0]
            caption = f"Figura {i+1}: {base_name.replace('_', ' ').title()}"
            
            # Adiciona a imagem com legenda
            success = add_image(doc, screenshot_path, caption)
            if success:
                print(f"Screenshot adicionado: {filename}")
            else:
                print(f"Falha ao adicionar screenshot: {filename}")
        else:
            print(f"Arquivo não encontrado: {screenshot_path}")
    
    # Adiciona a data e espaço para assinatura no final
    date_str = datetime.datetime.now().strftime("%d/%m/%Y")
    final_paragraph = doc.add_paragraph()
    final_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    final_paragraph.add_run(f"Dourados, {date_str} -- (Assinatura)")
    
    # Cria diretório de saída se não existir
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    # Salva o documento
    docx_path = f"{output_path}.docx"
    doc.save(docx_path)
    print(f"Documento DOCX salvo em: {docx_path}")
    
    # Converte para PDF se disponível
    if PDF_AVAILABLE:
        try:
            pdf_path = f"{output_path}.pdf"
            convert(docx_path, pdf_path)
            print(f"Documento PDF salvo em: {pdf_path}")
        except Exception as e:
            print(f"Erro ao converter para PDF: {str(e)}")
    
    return True

if __name__ == "__main__":
    args = parse_args()
    generate_screenshots_document(args.name, args.entrega, args.screenshots_list, args.output)
